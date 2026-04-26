from docx import Document
from docx.shared import Inches
from io import BytesIO
import requests


def create_weather_doc(city, current, daily_df):
    doc = Document()

    doc.add_heading(f"Real-Time Daily Weather Report - {city}", 0)

    doc.add_heading("Current Weather", level=1)
    doc.add_paragraph(f"City: {current['city']}, {current['country']}")
    doc.add_paragraph(f"Temperature: {current['temperature']} °C")
    doc.add_paragraph(f"Feels Like: {current['feels_like']} °C")
    doc.add_paragraph(f"Humidity: {current['humidity']} %")
    doc.add_paragraph(f"Pressure: {current['pressure']} hPa")
    doc.add_paragraph(f"Wind Speed: {current['wind']} m/s")
    doc.add_paragraph(f"Condition: {current['condition'].title()}")

    doc.add_paragraph("")
    doc.add_heading("Daily Forecast", level=1)

    for _, row in daily_df.iterrows():
        doc.add_heading(str(row["date"]), level=2)

        try:
            response = requests.get(row["icon_url"], timeout=10)
            image = BytesIO(response.content)
            doc.add_picture(image, width=Inches(1))
        except Exception:
            pass

        doc.add_paragraph(f"Condition: {row['condition'].title()}")
        doc.add_paragraph(f"Average Temperature: {row['temperature']} °C")
        doc.add_paragraph(f"Minimum Temperature: {row['temp_min']} °C")
        doc.add_paragraph(f"Maximum Temperature: {row['temp_max']} °C")
        doc.add_paragraph(f"Humidity: {int(row['humidity'])} %")
        doc.add_paragraph(f"Wind Speed: {row['wind']} m/s")

        doc.add_paragraph("----------------------------------------")

    file_path = f"{city}_real_time_daily_weather_report.docx"
    doc.save(file_path)

    return file_path